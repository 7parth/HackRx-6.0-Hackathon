import PyPDF2
import re
from typing import Dict, List, Tuple, Optional, Union
import logging
from pathlib import Path
import docx
import base64
import io

class DocumentRelevanceFilter:
    def __init__(self, min_relevance_threshold: float = 0.05):
        # Domain-specific keywords
        self.domain_keywords = {
        'insurance': [
            'policy', 'claim', 'coverage', 'premium', 'deductible', 'liability', 
            'underwriting', 'actuarial', 'risk assessment', 'insurance', 'policyholder',
            'beneficiary', 'indemnity', 'exclusion', 'riders', 'insurer', 'insured',
            'auto insurance', 'health insurance', 'life insurance', 'property insurance',
            'user manuel', 'SUPER SPLENDOR', 'Hero MotoCorp',

        ],

        'vehicle_manual': [
                'motorcycle', 'manual', 'user manual', 'maintenance', 'engine',
                'brakes', 'safety', 'specification', 'warranty', 'service',
                'fuel', 'clutch', 'transmission', 'electrical', 'chassis',
                'suspension', 'tyres', 'tires', 'battery', 'lubrication',
                'adjustment', 'troubleshooting', 'vehicle', 'Hero MotoCorp',
                'SUPER SPLENDOR', 'inspection', 'installation', 'operation',
                'starting', 'riding', 'spark plug', 'disc brake', 'drum brake',
                'tubeless', 'chain', 'oil', 'filter', 'cylinder', 'piston'
        ],
        'legal': [
            'contract', 'agreement', 'statute', 'regulation', 'compliance', 'litigation',
            'jurisdiction', 'defendant', 'plaintiff', 'arbitration', 'settlement',
            'legal', 'law', 'court', 'attorney', 'counsel', 'magistrate', 'judge',
            'legal notice', 'terms and conditions', 'privacy policy', 'disclaimer',

            # Enhanced keywords
            'constitution', 'constitutional', 'amendment', 'article', 'clause',
            'fundamental rights', 'directive principles', 'parliament', 'legislature',
            'judiciary', 'supreme court', 'high court', 'bill', 'act', 'legislation',
            'preamble', 'schedule', 'part', 'chapter', 'section', 'provision',
            'citizen', 'rights', 'duties', 'government', 'executive', 'judicial'
        ],
        'hr': [
            'employee', 'employment', 'payroll', 'benefits', 'performance', 'hiring',
            'termination', 'workplace', 'personnel', 'compensation', 'human resources',
            'training', 'onboarding', 'evaluation', 'disciplinary', 'workforce',
            'job description', 'salary', 'leave policy', 'performance review'
        ],
        'compliance': [
            'regulation', 'audit', 'compliance', 'governance', 'risk management',
            'internal controls', 'policy', 'procedure', 'standards', 'certification',
            'monitoring', 'reporting', 'oversight', 'framework', 'sox', 'gdpr',
            'regulatory compliance', 'audit report', 'control framework'
        ]
    }

        
        # Irrelevant content indicators
        self.irrelevant_indicators = [
            'physics', 'mathematics', 'scientific', 'principia', 'newton', 
            'recipe', 'cooking', 'food', 'entertainment', 'gaming', 'sports',
            'fiction', 'novel', 'story', 'biography',
            'academic research', 'thesis', 'dissertation', 'journal article' 
        ]
        
        self.min_relevance_threshold = min_relevance_threshold

    def extract_pdf_metadata(self, source: Union[str, bytes]) -> Dict:
        """Extract metadata from PDF file path or bytes"""
        file_obj = None
        try:
            if isinstance(source, bytes):
                pdf_stream = io.BytesIO(source)
                pdf_reader = PyPDF2.PdfReader(pdf_stream)
            else:  # File path
                file_obj = open(source, 'rb')
                pdf_reader = PyPDF2.PdfReader(file_obj)
            
            metadata = pdf_reader.metadata or {}
            first_page_text = ""
            if len(pdf_reader.pages) > 0:
                first_page_text = pdf_reader.pages[0].extract_text() or ""
                first_page_text = first_page_text[:1000]
            
            return {
                'title': str(metadata.get('/Title', '')).lower(),
                'subject': str(metadata.get('/Subject', '')).lower(),
                'keywords': str(metadata.get('/Keywords', '')).lower(),
                'creator': str(metadata.get('/Creator', '')).lower(),
                'producer': str(metadata.get('/Producer', '')).lower(),
                'author': str(metadata.get('/Author', '')).lower(),
                'first_page_sample': first_page_text.lower()[:500],
                'page_count': len(pdf_reader.pages),
                'file_type': 'pdf'
            }
        except Exception as e:
            logging.error(f"Error extracting PDF metadata: {e}")
            return {}
        finally:
            if file_obj:
                file_obj.close()

    def extract_docx_metadata(self, source: Union[str, bytes]) -> Dict:
        """Extract metadata from DOCX file path or bytes"""
        try:
            if isinstance(source, bytes):
                doc = docx.Document(io.BytesIO(source))
            else:  # File path
                doc = docx.Document(source) # type: ignore
                
            first_content = ""
            for paragraph in doc.paragraphs[:5]:
                first_content += paragraph.text + " "
                if len(first_content) > 500:
                    break
            
            core_props = doc.core_properties
            return {
                'title': str(core_props.title or '').lower(),
                'subject': str(core_props.subject or '').lower(),
                'keywords': str(core_props.keywords or '').lower(),
                'creator': str(core_props.author or '').lower(),
                'author': str(core_props.author or '').lower(),
                'first_page_sample': first_content.lower()[:500],
                'page_count': len(doc.paragraphs),
                'file_type': 'docx'
            }
        except Exception as e:
            logging.error(f"Error extracting DOCX metadata: {e}")
            return {}
        
    def calculate_relevance_score(self, metadata: Dict) -> Tuple[float, str, List[str]]:
        """Calculate relevance score with additional irrelevant indicators"""
        combined_text = ' '.join([
            metadata.get('title', ''),
            metadata.get('subject', ''),
            metadata.get('keywords', ''),
            metadata.get('creator', ''),
            metadata.get('author', ''),
            metadata.get('first_page_sample', '')
        ]).lower()
        
        if not combined_text.strip():
            return 0.0, "No metadata available", []
        
        # Check for irrelevant content first (immediate rejection)
        irrelevant_matches = [indicator for indicator in self.irrelevant_indicators 
                            if indicator in combined_text]
        
        if irrelevant_matches:
            matches_str = ', '.join(irrelevant_matches[:3])
            return 0.0, f"Contains irrelevant indicators: {matches_str}", irrelevant_matches
        
        # Calculate domain relevance
        domain_scores = {}
        domain_matches = {}
        
        for domain, keywords in self.domain_keywords.items():
            matches = [keyword for keyword in keywords if keyword in combined_text]
            domain_scores[domain] = len(matches)
            domain_matches[domain] = matches
        
        max_score = max(domain_scores.values()) if domain_scores else 0
        total_possible = max(len(keywords) for keywords in self.domain_keywords.values())
        
        relevance_score = max_score / total_possible if total_possible > 0 else 0
        best_domain = max(domain_scores.keys(), key=domain_scores.get) if max_score > 0 else "unknown" # type: ignore
        
        matched_keywords = domain_matches.get(best_domain, [])[:5]
        
        return relevance_score, f"Best match: {best_domain} ({max_score} keywords: {', '.join(matched_keywords)})", []

    def is_document_relevant(self, source: Union[str, bytes], file_type: str) -> Tuple[bool, str, Dict, List[str]]:
        """Check document relevance with enhanced return data"""
        try:
            # Extract metadata based on file type
            if file_type.lower() == 'pdf':
                metadata = self.extract_pdf_metadata(source)
            elif file_type.lower() in ['docx', 'doc']:
                metadata = self.extract_docx_metadata(source)
            else:
                return False, f"Unsupported file type: {file_type}", {}, []
            
            if not metadata:
                return False, "Failed to extract metadata", {}, []
            
            # Calculate relevance score with indicators
            relevance_score, reason, irrelevant_indicators = self.calculate_relevance_score(metadata)
            
            # Additional checks
            page_count = metadata.get('page_count', 0)
            if page_count < 1:
                return False, "Document appears to be empty", metadata, []
            
            # Determine relevance
            is_relevant = relevance_score >= self.min_relevance_threshold
            
            detailed_reason = f"Relevance score: {relevance_score:.3f} (threshold: {self.min_relevance_threshold}) - {reason}"
            
            return is_relevant, detailed_reason, metadata, irrelevant_indicators
            
        except Exception as e:
            logging.error(f"Error checking document relevance: {e}")
            return False, f"Error processing document: {str(e)}", {}, []

    
